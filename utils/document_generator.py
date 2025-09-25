import os
from datetime import datetime
from typing import Dict, Any
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import asyncio

from templates.document_templates import format_parking_dispute, format_housing_dispute

class DocumentGenerator:
    """Handles document generation for parking and housing disputes."""
    
    def __init__(self):
        self.output_dir = "output"
        # Create output directory if it doesn't exist
        os.makedirs(self.output_dir, exist_ok=True)
    
    async def generate_parking_dispute(self, data: Dict[str, Any]) -> str:
        """Generate a parking ticket dispute document."""
        # Format the document content
        content = format_parking_dispute(data)
        
        # Create Word document
        doc = Document()
        
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Add title
        title = doc.add_heading('PARKING CITATION DISPUTE', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add a subtitle
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.add_run(f"Citation Number: {data.get('ticket_number', 'N/A')}")
        subtitle_run.bold = True
        
        # Add document content
        paragraphs = content.split('\n\n')
        for paragraph_text in paragraphs:
            if paragraph_text.strip():
                if paragraph_text.strip().startswith('RE:') or paragraph_text.strip().startswith('VEHICLE INFORMATION:') or \
                   paragraph_text.strip().startswith('VIOLATION ALLEGED:') or paragraph_text.strip().startswith('GROUNDS FOR DISPUTE:') or \
                   paragraph_text.strip().startswith('SUPPORTING EVIDENCE:') or paragraph_text.strip().startswith('LEGAL BASIS FOR DISMISSAL:') or \
                   paragraph_text.strip().startswith('CONCLUSION:') or paragraph_text.strip().startswith('ATTACHMENTS:'):
                    # Headers in bold
                    p = doc.add_paragraph()
                    p.add_run(paragraph_text.strip()).bold = True
                else:
                    # Regular paragraphs
                    doc.add_paragraph(paragraph_text.strip())
        
        # Generate filename with timestamp
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"parking_dispute_{timestamp}.docx"
        filepath = os.path.join(self.output_dir, filename)
        
        # Save document
        doc.save(filepath)
        
        return os.path.abspath(filepath)
    
    async def generate_housing_dispute(self, data: Dict[str, Any]) -> str:
        """Generate a housing dispute document."""
        # Format the document content
        content = format_housing_dispute(data)
        
        # Create Word document
        doc = Document()
        
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Add title
        title = doc.add_heading('FORMAL HOUSING COMPLAINT', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add a subtitle
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        property_info = data.get("property_info", "")
        property_address = "N/A"
        for line in property_info.split('\n'):
            if line.lower().startswith("property address:") or line.lower().startswith("address:"):
                property_address = line.split(":", 1)[1].strip()
                break
        subtitle_run = subtitle.add_run(f"Property: {property_address}")
        subtitle_run.bold = True
        
        # Add document content
        paragraphs = content.split('\n\n')
        for paragraph_text in paragraphs:
            if paragraph_text.strip():
                if paragraph_text.strip().startswith('RE:') or paragraph_text.strip().startswith('PROPERTY INFORMATION:') or \
                   paragraph_text.strip().startswith('ISSUE DESCRIPTION:') or paragraph_text.strip().startswith('TIMELINE OF EVENTS:') or \
                   paragraph_text.strip().startswith('PREVIOUS ATTEMPTS AT RESOLUTION:') or paragraph_text.strip().startswith('IMPACT ON HABITABILITY:') or \
                   paragraph_text.strip().startswith('REQUESTED RESOLUTION:') or paragraph_text.strip().startswith('LEGAL OBLIGATIONS:') or \
                   paragraph_text.strip().startswith('SUPPORTING DOCUMENTATION:') or paragraph_text.strip().startswith('TIMELINE FOR RESPONSE:') or \
                   paragraph_text.strip().startswith('NEXT STEPS:') or paragraph_text.strip().startswith('COPIES SENT TO:') or \
                   paragraph_text.strip().startswith('ATTACHMENTS:'):
                    # Headers in bold
                    p = doc.add_paragraph()
                    p.add_run(paragraph_text.strip()).bold = True
                else:
                    # Regular paragraphs
                    doc.add_paragraph(paragraph_text.strip())
        
        # Generate filename with timestamp
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"housing_dispute_{timestamp}.docx"
        filepath = os.path.join(self.output_dir, filename)
        
        # Save document
        doc.save(filepath)
        
        return os.path.abspath(filepath)
    
    def create_simple_text_document(self, content: str, filename: str) -> str:
        """Create a simple text document as fallback."""
        filepath = os.path.join(self.output_dir, f"{filename}.txt")
        
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(content)
        
        return os.path.abspath(filepath)